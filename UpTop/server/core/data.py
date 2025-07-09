import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create the document
doc = Document()

# Add title
title = doc.add_heading("Academic Plan", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add content
text = (
    "My name is Aziza Karimberganova Quvandiq qizi, and I am currently studying History at Mamun University. "
    "I was accepted in 2024 and now I am a first-year student at the Faculty of Social, Humanitarian and Exact Sciences.\n\n"
    "Through the student exchange program, I would like to experience studying in Korea, a country with rich culture and unique traditions. "
    "I am especially interested in learning the Korean language and understanding more about Korean history and lifestyle. "
    "I believe that living and studying in Korea will give me the chance to not only improve my language skills but also develop a more global view of the world.\n\n"
    "As a history student, I love exploring how different cultures have developed and how people live in different parts of the world. "
    "I think Korea is one of the best places to learn both old and modern culture at the same time.\n\n"
    "In the future, I want to become a teacher and help young students, especially children, to enjoy learning. "
    "I believe this exchange experience will help me grow both as a student and a future educator. "
    "I hope to share what I learn in Korea with others back home and maybe even help build cultural understanding between our countries.\n\n"
    "Thank you for the opportunity.\n\n"
    "Sincerely,\n"
    "Aziza Karimberganova Quvandiq qizi"
)

doc.add_paragraph(text)

# Correct file path
desktop_path = os.path.expanduser("~/Desktop/Academic_Plan_Aziza_Karimberganova.docx")
doc.save(desktop_path)
